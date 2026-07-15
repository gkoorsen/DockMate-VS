#!/usr/bin/env python3
"""
Generate the student WSL setup guide as a Word (.docx) document.

Produces docs/SETUP_WSL.docx: a step-by-step guide that starts from the Ubuntu
(WSL) shell — it assumes the student already has WSL + Ubuntu installed and is
at the bash prompt — and takes them through installing the tools, cloning the
repository, creating the conda environment, and launching the GUI.

Usage:
    python scripts/make_wsl_setup_doc.py [--output docs/SETUP_WSL.docx]
"""

import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

CODE_FONT = "Consolas"
CODE_SHADE = "F2F2F2"   # light grey
NOTE_SHADE = "FFF6D5"   # light amber
ACCENT = RGBColor(0x1F, 0x4E, 0x79)


def _shade(paragraph, fill_hex):
    """Apply a solid background fill to a paragraph."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    paragraph.paragraph_format.element.get_or_add_pPr().append(shd)


def add_code(doc, text):
    """Add a shaded, monospaced code block (one paragraph, preserving lines)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(10)
    lines = text.strip("\n").split("\n")
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = CODE_FONT
        run.font.size = Pt(10)
        if i < len(lines) - 1:
            run.add_break()
    _shade(p, CODE_SHADE)
    return p


def add_note(doc, text):
    """Add a shaded 'note' callout paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Note:  ")
    run.bold = True
    p.add_run(text)
    _shade(p, NOTE_SHADE)
    return p


def add_steps(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_shell_label(doc, text="Run this in the Ubuntu (WSL) terminal:"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p.paragraph_format.space_after = Pt(2)


def build(doc):
    # --- Styling defaults ---
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = doc.add_heading("Setting up the Docking GUI in WSL (Ubuntu)", level=0)
    title.runs[0].font.color.rgb = ACCENT

    intro = doc.add_paragraph()
    intro.add_run(
        "This guide assumes you already have WSL with Ubuntu installed and you "
        "are looking at the Ubuntu shell prompt. "
    )
    intro.add_run("Every command below runs inside the Ubuntu (WSL) terminal.").bold = True
    doc.add_paragraph(
        "Follow the steps in order. Most of the time is spent waiting for "
        "downloads (allow ~30 minutes and about 15 GB of free disk space)."
    )

    # --- What you'll end up with ---
    doc.add_heading("What you'll end up with", level=1)
    add_bullets(doc, [
        "A conda environment called docking_platform_gui with Python, RDKit, "
        "the docking engines, and every other tool the GUI needs.",
        "The GUI window opening on your Windows desktop.",
    ])

    # --- Step 1 ---
    doc.add_heading("Step 1 — Update Ubuntu and install the basics", level=1)
    add_shell_label(doc)
    add_code(doc, "sudo apt update && sudo apt upgrade -y\nsudo apt install -y git wget")
    doc.add_paragraph(
        "When you run a sudo command it asks for your Ubuntu password. It is "
        "normal that nothing appears on screen while you type it — just type it "
        "and press Enter."
    )

    # --- Step 2 ---
    doc.add_heading("Step 2 — Quick check that a window can open (optional)", level=1)
    doc.add_paragraph(
        "The GUI needs to open a window. On Windows 11 and Windows 10 22H2+ this "
        "works automatically. To confirm:"
    )
    add_shell_label(doc)
    add_code(doc, "sudo apt install -y x11-apps\nxeyes")
    doc.add_paragraph(
        "A small pair of googly eyes should appear on your desktop. If it does, "
        "graphics work — close it and continue. If nothing appears, see "
        "Troubleshooting at the end."
    )

    # --- Step 3 ---
    doc.add_heading("Step 3 — Install Miniconda", level=1)
    doc.add_paragraph(
        "Conda installs Python and the compiled docking tools (smina, Open "
        "Babel, fpocket, OpenMM, …) that plain pip cannot."
    )
    add_shell_label(doc)
    add_code(doc,
        "wget https://repo.anaconda.com/miniconda/Miniconda3-latest-Linux-x86_64.sh -O ~/miniconda.sh\n"
        "bash ~/miniconda.sh -b -p $HOME/miniconda3\n"
        "~/miniconda3/bin/conda init bash")
    doc.add_paragraph(
        "Close the Ubuntu terminal and open it again so conda activates. You "
        "should now see (base) at the start of the prompt. Check it works:"
    )
    add_code(doc, "conda --version")

    # --- Step 4 ---
    doc.add_heading("Step 4 — Get the code", level=1)
    add_note(doc,
        "Clone into your Linux home folder (~) as shown. Do NOT put it under "
        "/mnt/c/... (your Windows C: drive) — it will be much slower and can "
        "cause permission errors.")
    add_shell_label(doc)
    add_code(doc,
        "mkdir -p ~/projects\n"
        "cd ~/projects\n"
        "git clone https://github.com/gkoorsen/docking_platform_gui.git\n"
        "cd docking_platform_gui")
    doc.add_paragraph("Then switch to the branch that has the Linux setup fixes:")
    add_code(doc, "git checkout claude/linux-install-fixes")
    add_note(doc,
        "Once that branch has been merged into master you can skip this "
        "git checkout and just use master.")

    doc.add_heading("If the repository is private", level=2)
    doc.add_paragraph(
        "The git clone will ask you to sign in. The simplest way is the GitHub "
        "CLI, which stores your login for all future git commands:"
    )
    add_code(doc, "sudo apt install -y gh\ngh auth login")
    doc.add_paragraph(
        "Choose GitHub.com → HTTPS → \"Login with a web browser\" and follow the "
        "prompts, then run the git clone command again. (Alternatively, when git "
        "asks for a password, paste a GitHub Personal Access Token instead of "
        "your account password.)"
    )

    # --- Step 5 ---
    doc.add_heading("Step 5 — Create the environment", level=1)
    doc.add_paragraph("Still inside ~/projects/docking_platform_gui:")
    add_code(doc, "conda env create -f environment.yml")
    doc.add_paragraph(
        "This installs everything, including the package itself. It can take "
        "10–20 minutes, and it is normal for the \"Solving environment\" step to "
        "sit for a while. When it finishes, activate the environment:"
    )
    add_code(doc, "conda activate docking_platform_gui")
    p = doc.add_paragraph()
    p.add_run("Your prompt should now start with (docking_platform_gui). ")
    p.add_run(
        "You must run this conda activate command every time you open a new "
        "Ubuntu terminal to use the GUI."
    ).bold = True

    # --- Step 6 ---
    doc.add_heading("Step 6 — Launch the GUI", level=1)
    add_code(doc, "redock-gui")
    doc.add_paragraph(
        "The Docking GUI window should open on your Windows desktop. "
        "(Equivalently: python scripts/launch_redock_analysis_gui.py.)"
    )

    # --- Using it afterwards ---
    doc.add_heading("Using it again later", level=1)
    doc.add_paragraph("Each time you want to run the GUI, open the Ubuntu terminal and run:")
    add_code(doc,
        "cd ~/projects/docking_platform_gui\n"
        "conda activate docking_platform_gui\n"
        "redock-gui")
    doc.add_paragraph("To update to the latest code later:")
    add_code(doc,
        "cd ~/projects/docking_platform_gui\n"
        "git pull\n"
        "conda env update -f environment.yml --prune   # only if environment.yml changed")

    # --- Troubleshooting ---
    doc.add_heading("Troubleshooting", level=1)
    tips = [
        ("The GUI window never appears / xeyes didn't work.",
         "In Windows PowerShell run `wsl --update`, then `wsl --shutdown`, "
         "reopen Ubuntu and try again. Make sure Windows is fully updated and "
         "you are on Windows 11 or Windows 10 22H2+. Inside Ubuntu, `echo "
         "$DISPLAY` should print something like :0 — if it's empty, the display "
         "isn't running."),
        ("conda: command not found.",
         "You skipped reopening the terminal after conda init. Close and reopen "
         "Ubuntu, or run `source ~/.bashrc`."),
        ("conda env create is very slow or fails to \"solve\".",
         "Update conda first: `conda update -n base -c defaults conda`. If one "
         "specific tool refuses to install, open environment.yml, put a # in "
         "front of that one line, and re-run — the GUI will still start."),
        ("redock-gui: command not found.",
         "The environment isn't active. Run `conda activate docking_platform_gui` "
         "— the (docking_platform_gui) prefix must be showing."),
        ("A tool like smina or obabel is reported missing.",
         "Make sure the environment is active; these tools live inside the "
         "docking_platform_gui environment only."),
        ("Everything feels slow.",
         "Make sure you cloned into ~/projects (Linux side), not /mnt/c/... "
         "(Windows side)."),
    ]
    for heading, body in tips:
        p = doc.add_paragraph()
        p.add_run(heading).bold = True
        doc.add_paragraph(body)

    # --- Quick reference ---
    doc.add_heading("Quick reference (all commands in order)", level=1)
    add_shell_label(doc)
    add_code(doc,
        "sudo apt update && sudo apt upgrade -y\n"
        "sudo apt install -y git wget x11-apps\n"
        "xeyes                                   # test graphics, then close it\n"
        "\n"
        "wget https://repo.anaconda.com/miniconda/Miniconda3-latest-Linux-x86_64.sh -O ~/miniconda.sh\n"
        "bash ~/miniconda.sh -b -p $HOME/miniconda3\n"
        "~/miniconda3/bin/conda init bash\n"
        "# close and reopen the Ubuntu terminal here\n"
        "\n"
        "mkdir -p ~/projects && cd ~/projects\n"
        "git clone https://github.com/gkoorsen/docking_platform_gui.git\n"
        "cd docking_platform_gui\n"
        "git checkout claude/linux-install-fixes\n"
        "\n"
        "conda env create -f environment.yml\n"
        "conda activate docking_platform_gui\n"
        "redock-gui")


def main():
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", default="docs/SETUP_WSL.docx")
    args = parser.parse_args()

    doc = Document()
    build(doc)
    out = Path(args.output)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"Wrote {out.resolve()} ({out.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
